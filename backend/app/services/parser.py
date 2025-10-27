"""
Document parser service for various file formats
"""
from docx import Document as DocxDocument
from pypdf import PdfReader
from typing import Dict, Tuple
from pathlib import Path
import logging
import asyncio

from app.models.document import DocumentType

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse various document formats"""
    
    @staticmethod
    async def parse_file(file_path: Path, doc_type: DocumentType) -> str:
        """
        Parse document based on type (async wrapper)
        
        Args:
            file_path: Path to document
            doc_type: DocumentType enum
            
        Returns:
            Extracted text content
        """
        if doc_type == DocumentType.WORD:
            text, _ = DocumentParser.parse_docx(str(file_path))
            return text
        elif doc_type == DocumentType.PDF:
            text, _ = DocumentParser.parse_pdf(str(file_path))
            return text
        elif doc_type == DocumentType.TEXT:
            return await DocumentParser.parse_text_async(str(file_path))
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, Dict]:
        """
        Parse Word document
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text
            full_text = []
            sections = []
            current_section = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'header': text,
                        'level': para.style.name,
                        'content': []
                    }
                elif current_section is not None:
                    current_section['content'].append(text)
                
                full_text.append(text)
            
            # Add last section
            if current_section:
                sections.append(current_section)
            
            text_content = '\n\n'.join(full_text)
            
            # Extract metadata
            metadata = {
                'paragraphs_count': len(doc.paragraphs),
                'sections_count': len(sections),
                'sections': sections,
                'core_properties': {
                    'author': doc.core_properties.author or '',
                    'title': doc.core_properties.title or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                }
            }
            
            logger.info(f"Parsed DOCX: {len(full_text)} paragraphs, {len(sections)} sections")
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, Dict]:
        """
        Parse PDF document
        
        Args:
            file_path: Path to .pdf file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                
                full_text = '\n\n'.join(text_content)
                
                # Extract metadata
                metadata = {
                    'pages_count': len(pdf_reader.pages),
                    'info': {}
                }
                
                if pdf_reader.metadata:
                    metadata['info'] = {
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', '')
                    }
                
                logger.info(f"Parsed PDF: {len(pdf_reader.pages)} pages")
                return full_text, metadata
                
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    @staticmethod
    def parse_text(file_path: str) -> Tuple[str, Dict]:
        """
        Parse plain text file (sync version)
        
        Args:
            file_path: Path to text file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Basic metadata
            lines = text_content.split('\n')
            metadata = {
                'lines_count': len(lines),
                'characters_count': len(text_content),
                'words_count': len(text_content.split())
            }
            
            logger.info(f"Parsed text file: {len(lines)} lines")
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing text file: {str(e)}")
            raise
    
    @staticmethod
    async def parse_text_async(file_path: str) -> str:
        """
        Parse plain text file (async version)

        Args:
            file_path: Path to text file

        Returns:
            Text content
        """
        try:
            # Use asyncio to run file I/O in executor
            loop = asyncio.get_event_loop()
            text_content = await loop.run_in_executor(
                None,
                lambda: open(file_path, 'r', encoding='utf-8').read()
            )

            logger.info(f"Parsed text file async: {len(text_content)} chars")
            return text_content

        except Exception as e:
            logger.error(f"Error parsing text file async: {str(e)}")
            raise
    
    @staticmethod
    def parse_document(file_path: str, doc_type: str) -> Tuple[str, Dict]:
        """
        Parse document based on type string (for backward compatibility)
        
        Args:
            file_path: Path to document
            doc_type: Type of document as string (word, pdf, text)
            
        Returns:
            Tuple of (text_content, metadata)
        """
        if doc_type == "word":
            return DocumentParser.parse_docx(file_path)
        elif doc_type == "pdf":
            return DocumentParser.parse_pdf(file_path)
        elif doc_type in ["text", "markdown"]:
            return DocumentParser.parse_text(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    @staticmethod
    def extract_metadata_only(file_path: str, doc_type: DocumentType) -> Dict:
        """
        Extract only metadata without full text
        
        Args:
            file_path: Path to document
            doc_type: DocumentType enum
            
        Returns:
            Metadata dict
        """
        try:
            if doc_type == DocumentType.WORD:
                _, metadata = DocumentParser.parse_docx(file_path)
                return metadata
            elif doc_type == DocumentType.PDF:
                _, metadata = DocumentParser.parse_pdf(file_path)
                return metadata
            elif doc_type == DocumentType.TEXT:
                _, metadata = DocumentParser.parse_text(file_path)
                return metadata
            else:
                return {}
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}


# Global instance
document_parser = DocumentParser()